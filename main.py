from keep_alive import keep_alive
keep_alive()
# -*- coding: utf-8 -*-
import os
import sqlite3
import telebot
import PyPDF2
# import google.generativeai as genai # Removed Gemini
import requests # Added for API calls
import json # Added for API calls
import matplotlib
matplotlib.use('Agg') # Use Agg backend for non-interactive plotting
import matplotlib.pyplot as plt
import numpy as np
import io
import sympy as sp
import re
from PIL import Image # Required by pytesseract for some image formats
import tempfile
from telebot import types
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ==============================================================================
# Configuration & Setup
# ==============================================================================

# --- Paths for Tesseract and Poppler ---

# المسار الصحيح لتثبيت tesseract في Docker
pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

# المسار الصحيح لأدوات poppler في Docker
poppler_path = "/usr/bin/"
# هذا هو المسار الذي أكدته
# For Linux; adjust for other OS if needed

# --- API Keys ---
# تأكد من أن هذا هو مفتاح OpenRouter API الخاص بك
OPENROUTER_API_KEY = "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"
TELEGRAM_TOKEN = "8076364578:AAH9KZHPBGT7NZt_Pu4-GPdhb-apAHoU7O8" # Replace with your Telegram Bot Token

# --- OpenRouter API Configuration ---
# مثال: "deepseek/deepseek-chat", "openai/gpt-4o", "anthropic/claude-3-opus"
# تأكد من استخدام اسم النموذج الصحيح المتوفر على OpenRouter
OPENROUTER_MODEL_NAME = "deepseek/deepseek-chat"
OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
# يمكنك تعديل هذه القيم حسب الحاجة
YOUR_SITE_URL = "https://github.com/YoussifMK/Sooma-Bot" # أو أي رابط لموقعك/مشروعك
YOUR_APP_NAME = "Sooma_Telegram_Bot" # اسم تطبيقك

GENERATION_CONFIG = {
    "temperature": 1.0,
    "top_p": 0.95,
    # "top_k": 40, # top_k is often not used with top_p in OpenAI-like APIs
    "max_output_tokens": 8192, # Corresponds to max_tokens
    # "response_mime_type": "text/plain", # Not applicable for JSON response
}

model_ready = False
if OPENROUTER_API_KEY and OPENROUTER_API_KEY != "sk-or-v1-YOUR_OPENROUTER_API_KEY_HERE": # Check against a generic placeholder
    model_ready = True
else:
    print("Warning: Please set your actual OpenRouter API key.")

# --- Initialize Telegram Bot ---
if TELEGRAM_TOKEN == "YOUR_TELEGRAM_BOT_TOKEN":
    print("Warning: Please replace 'YOUR_TELEGRAM_BOT_TOKEN' with your actual Telegram Bot token.")
bot = telebot.TeleBot(TELEGRAM_TOKEN, parse_mode=None)

# --- Active Sessions Storage ---
# Stores user-specific data like extracted PDF text or active chat history
# Structure: active_sessions[user_id] = {"status": "...", "extracted_text": "...", "chat_history": [...]}
active_sessions = {}

# --- User Identification ---
correct_yusuf_full_name = "يوسف محمد كمال أبو رومية"
yusuf_aliases = ["يوسف محمد", "يوسف محمد كمال", "يوسف محمد كمال أبو رومية", "يوسف أبو رومية", "يوسف"]
correct_assim_full_name = "عاصم ايهاب ثابت دردير"
assim_aliases = ["عاصم ايهاب", "عاصم ايهاب ثابت", "عاصم ايهاب ثابت دردير", "عاصم"]
correct_basmala_full_name = "بسملة محمد محمد مبروك"
basmala_aliases = ["بسملة محمد", "بسملة محمد محمد", "بسملة محمد مبروك", "بسملة"]

user_names_identified = {} # Stores identified user names: user_id -> "يوسف" / "عاصم" / "بسملة"
greeting_words = ["مرحبا", "السلام عليكم", "اهلا", "أهلا", "هاي", "hello"]

# ==============================================================================
# Database Functions (SQLite)
# ==============================================================================
DATABASE_NAME = 'users.db'

def setup_database():
    """Creates the users table if it doesn't exist."""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS users (
        user_id INTEGER PRIMARY KEY,
        username TEXT,
        is_special INTEGER DEFAULT 0 
    )
    ''') # is_special: 0=normal, 1=special, 2=super_special
    conn.commit()
    conn.close()

def add_user(user_id, username, is_special=0):
    """Adds or replaces a user in the database."""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    cursor.execute(
        'INSERT OR REPLACE INTO users (user_id, username, is_special) VALUES (?, ?, ?)',
        (user_id, username, is_special)
    )
    conn.commit()
    conn.close()

def is_user_special(user_id):
    """
    Checks the special status of a user.
    Returns: 0 (normal), 1 (special), 2 (super_special), or 0 if user not found.
    """
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    cursor.execute('SELECT is_special FROM users WHERE user_id = ?', (user_id,))
    result = cursor.fetchone()
    conn.close()
    return result[0] if result else 0

def update_user_username_if_changed(user_id, current_username):
    """Updates the username in the DB if it has changed in Telegram."""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    cursor.execute('SELECT username FROM users WHERE user_id = ?', (user_id,))
    result = cursor.fetchone()
    if result:
        stored_username = result[0]
        if stored_username != current_username:
            cursor.execute('UPDATE users SET username = ? WHERE user_id = ?', (current_username, user_id))
            conn.commit()
            print(f"Updated username for user_id {user_id} from '{stored_username}' to '{current_username}'")
    else:
        add_user(user_id, current_username, 0)
        print(f"New user added: User ID: {user_id}, Username: {current_username}")
    conn.close()

# ==============================================================================
# OpenRouter API Interaction Function
# ==============================================================================
def send_message_to_openrouter(history, config):
    """
    Sends a message history to the OpenRouter API and gets a response.
    history: list of messages in OpenAI format, e.g., [{"role": "user", "content": "Hello"}]
    config: dictionary with generation parameters like temperature, max_tokens, top_p
    """
    if not model_ready:
        return "Error: OpenRouter API not configured properly."

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "HTTP-Referer": YOUR_SITE_URL, # Recommended by OpenRouter
        "X-Title": YOUR_APP_NAME,      # Recommended by OpenRouter
    }

    payload = {
        "model": OPENROUTER_MODEL_NAME,
        "messages": history,
        "temperature": config.get("temperature", 1.0),
        "max_tokens": config.get("max_output_tokens", 8192), # OpenRouter uses "max_tokens"
        "top_p": config.get("top_p", 0.95),
        # "stream": False # Keep it non-streaming
    }
    # Filter out any None values from the payload
    payload = {k: v for k, v in payload.items() if v is not None}

    try:
        response = requests.post(OPENROUTER_API_URL, headers=headers, json=payload, timeout=180) # Increased timeout
        response.raise_for_status()  # Raises an HTTPError for bad responses (4XX or 5XX)

        response_data = response.json()

        if "choices" in response_data and response_data["choices"]:
            message_content = response_data["choices"][0].get("message", {}).get("content")
            if message_content:
                return message_content.strip()
            else:
                print(f"OpenRouter API response missing content: {response_data}")
                return "Error: AI response was empty or malformed."
        else:
            print(f"OpenRouter API response missing choices: {response_data}")
            return "Error: AI did not provide a valid response choice."

    except requests.exceptions.Timeout:
        print(f"Timeout error calling OpenRouter API.")
        return "Error: The request to the AI timed out."
    except requests.exceptions.RequestException as e:
        print(f"Error calling OpenRouter API: {e}")
        error_response_text = ""
        if hasattr(e, 'response') and e.response is not None:
            try:
                error_response_text = e.response.text
            except Exception:
                pass # Ignore if can't get text
        print(f"Response status: {e.response.status_code if hasattr(e, 'response') else 'N/A'}")
        print(f"Response content: {error_response_text}")
        return f"Error communicating with AI: {str(e)}. Check server logs for more details."
    except (KeyError, IndexError, json.JSONDecodeError) as e:
        print(f"Error parsing OpenRouter API response: {e}")
        # response_text_for_debug = response.text if 'response' in locals() and response is not None else "No response object"
        # print(f"Response content that caused parsing error: {response_text_for_debug}")
        return "Error parsing AI response. The format was unexpected."
    except Exception as e: # Catch any other unexpected errors
        print(f"An unexpected error occurred with OpenRouter API: {e}")
        return f"An unexpected error occurred while contacting the AI: {str(e)}"

# ==============================================================================
# PDF Processing Functions
# ==============================================================================
def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a PDF file.
    Tries PyPDF2 first, then falls back to OCR (Tesseract) if no text is found.
    """
    text = ""
    try:
        with open(pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            if reader.is_encrypted:
                try:
                    reader.decrypt('') 
                except Exception as e:
                    print(f"PDF is encrypted and could not be decrypted: {e}")

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF with PyPDF2: {e}")

    if text.strip():
        print("Text extracted using PyPDF2.")
        return text.strip()

    print("PyPDF2 extracted no text or failed. Attempting OCR...")
    try:
        images = convert_from_path(pdf_path, dpi=300, poppler_path=poppler_path)
        ocr_text = ""
        for i, img in enumerate(images):
            print(f"OCR processing page {i+1}/{len(images)}")
            ocr_text += pytesseract.image_to_string(img, lang='ara+eng') + "\n"

        if ocr_text.strip():
            print("Text extracted using OCR.")
            return ocr_text.strip()
        else:
            print("OCR did not find any text.")
            return ""
    except Exception as e:
        print(f"Error during OCR processing: {e}")
        print("Make sure Tesseract is installed and configured, and Poppler path is correct.")
        return ""

# ==============================================================================
# DOCX Formatting and Generation Functions
# ==============================================================================
def clean_markdown_for_docx(text):
    """Cleans text from unwanted markdown for DOCX generation."""
    text = re.sub(r'^\s*#+\s*', '', text, flags=re.MULTILINE)
    text = text.replace('**', '').replace('*', '')
    text = text.replace('`', '')
    return text

def process_text_with_openrouter_for_docx(text_content, format_caption):
    """
    Sends text and formatting instructions to OpenRouter to get structured output for DOCX.
    """
    if not model_ready:
        return "Error: AI model not initialized."

    prompt_text = f"""
هذا نص مستخرج من ملف PDF. المطلوب منك تنسيقه في ملف Word بناءً على الطلب التالي:
"{format_caption}"

الرجاء اتباع التعليمات التالية بدقة:

- قم بتنسيق العناوين بخط عريض (bold) واستخدم حجم أكبر قليلاً من باقي النص (size). مثال: <bold><size=14>هذا عنوان</size=14></bold>
- لا تستخدم الخط المائل (italic) إلا إذا طُلب صراحةً. إذا طُلب، استخدم <italic>النص المائل</italic>.
- إذا طلب المستخدم تلوين أي جزء من النص (مثلاً "اجعل العنوان باللون الأحمر الداكن")، الرجاء تحويل اسم اللون إلى كود Hex تلقائيًا، واستخدمه مثل: <color=#8B0000>النص الملون</color>.
- أنشئ قوائم نقطية باستخدام '-' في بداية كل عنصر، أو مرقمة باستخدام '1.'، '2.' في بداية كل عنصر. لا تستخدم وسوم خاصة للقوائم.
- ضع المعادلات في سطر منفصل، ويفضّل توسيطها باستخدام: <align=center>المعادلة</align=center>.
- إذا طُلب إنشاء جدول، استخدم الشكل التالي بالضبط (كل وسم في سطر جديد):
<table_start>
<row>
<cell>العنوان الأول</cell>
<cell>العنوان الثاني</cell>
</row>
<row>
<cell>بيانات الخلية 1</cell>
<cell>بيانات الخلية 2</cell>
</row>
<table_end>

ملاحظات مهمة:
- استخدم فقط الوسوم التالية للتنسيق المباشر: <bold>...</bold>, <italic>...</italic>, <underline>...</underline>, <size=VALUE>...</size=VALUE> (حيث VALUE هو رقم حجم الخط بالنقاط، مثال <size=12>), <color=#HEXCODE>...</color=#HEXCODE>, <align=POSITION>...</align=POSITION> (حيث POSITION يمكن أن تكون center, right, left).
- تأكد من إغلاق جميع الوسوم بشكل صحيح. مثال: <bold>نص عريض</bold>.
- لا تستخدم رموز تنسيق Markdown مثل * أو # أو ` في إجابتك النهائية.
- النص الأصلي قد يحتوي على تنسيقات Markdown، تجاهلها وركز على تطبيق الوسوم المذكورة أعلاه بناءً على طلب المستخدم.

النص الأصلي من الـ PDF:
{text_content}
"""
    docx_history = [
        {"role": "system", "content": "أنت مساعد متخصص في تنسيق النصوص لملفات DOCX باستخدام وسوم خاصة."},
        {"role": "user", "content": prompt_text}
    ]

    response_text = send_message_to_openrouter(docx_history, GENERATION_CONFIG)
    return response_text


def parse_formatted_text_for_docx(formatted_text_from_ai):
    paragraphs_data = []
    table_data = []
    in_table = False
    current_row_cells = []

    lines = formatted_text_from_ai.replace('\r\n', '\n').split('\n')

    for line_raw in lines:
        line = line_raw.strip()

        if line == '<table_start>':
            in_table = True
            table_data = [] 
            continue
        elif line == '<table_end>':
            in_table = False
            if current_row_cells: 
                print("Warning: <table_end> found with pending cells in current_row_cells.")
                current_row_cells = []
            continue

        if in_table:
            if line == '<row>':
                current_row_cells = []
            elif line == '</row>':
                if current_row_cells:
                    table_data.append(list(current_row_cells)) 
                    current_row_cells = []
            elif line.startswith('<cell>') and line.endswith('</cell>'):
                cell_text = line[len('<cell>'):-len('</cell>')].strip()
                current_row_cells.append(cell_text)
            elif line: 
                print(f"Warning: Unexpected line inside table block: '{line}'")
        else: 
            if not line: 
                continue

            text_content = line
            style = {}

            if '<bold>' in text_content and '</bold>' in text_content:
                style['bold'] = True
                text_content = re.sub(r'</?bold>', '', text_content)
            if '<italic>' in text_content and '</italic>' in text_content:
                style['italic'] = True
                text_content = re.sub(r'</?italic>', '', text_content)
            if '<underline>' in text_content and '</underline>' in text_content:
                style['underline'] = True
                text_content = re.sub(r'</?underline>', '', text_content)

            color_match = re.search(r'<color=([^>]+)>(.*?)</color(?:=\1)?>', text_content, re.IGNORECASE)
            if not color_match:
                color_match = re.search(r'<color=([^>]+)>(.*?)</color>', text_content, re.IGNORECASE)
            if color_match:
                style['color'] = color_match.group(1).strip()
                text_content = re.sub(r'<color=[^>]+>(.*?)</color(?:=[^>]+)?>', r'\1', text_content, flags=re.IGNORECASE)

            size_match = re.search(r'<size=([0-9]+)>(.*?)</size(?:=\1)?>', text_content, re.IGNORECASE)
            if not size_match:
                size_match = re.search(r'<size=([0-9]+)>(.*?)</size>', text_content, re.IGNORECASE)
            if size_match:
                style['size'] = int(size_match.group(1))
                text_content = re.sub(r'<size=[0-9]+>(.*?)</size(?:=[0-9]+)?>', r'\1', text_content, flags=re.IGNORECASE)

            align_match = re.search(r'<align=([^>]+)>(.*?)</align(?:=\1)?>', text_content, re.IGNORECASE)
            if not align_match:
                align_match = re.search(r'<align=([^>]+)>(.*?)</align>', text_content, re.IGNORECASE)
            if align_match:
                align_str = align_match.group(1).lower().strip()
                if align_str == 'center':
                    style['alignment'] = WD_ALIGN_PARAGRAPH.CENTER
                elif align_str == 'right':
                    style['alignment'] = WD_ALIGN_PARAGRAPH.RIGHT
                elif align_str == 'left':
                    style['alignment'] = WD_ALIGN_PARAGRAPH.LEFT
                text_content = re.sub(r'<align=[^>]+>(.*?)</align(?:=[^>]+)?>', r'\1', text_content, flags=re.IGNORECASE)

            text_content_cleaned = re.sub(r'</?(?:bold|italic|underline|color|size|align)[^>]*>', '', text_content).strip()

            if text_content_cleaned: 
                paragraphs_data.append({'text': text_content_cleaned, 'style': style})
            elif style and not text_content_cleaned: 
                print(f"Warning: Line with only style tags and no content: '{line_raw}'")
    return paragraphs_data, table_data

def save_to_docx(paragraphs_data, table_data, output_path):
    doc = Document()
    for item in paragraphs_data:
        text = item['text']
        style = item.get('style', {})
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)

        if style.get('bold'): run.font.bold = True
        if style.get('italic'): run.font.italic = True
        if style.get('underline'): run.font.underline = True
        if 'size' in style:
            try: run.font.size = Pt(int(style['size']))
            except ValueError: print(f"Warning: Invalid font size '{style['size']}' for text '{text}'")
        if 'color' in style:
            try:
                color_str = style['color'].lstrip('#')
                if len(color_str) == 6: 
                    r,g,b = int(color_str[0:2],16), int(color_str[2:4],16), int(color_str[4:6],16)
                    run.font.color.rgb = RGBColor(r, g, b)
                elif len(color_str) == 3: 
                    r,g,b = int(color_str[0]*2,16), int(color_str[1]*2,16), int(color_str[2]*2,16)
                    run.font.color.rgb = RGBColor(r, g, b)   
                else: print(f"Warning: Invalid color hex '{style['color']}' for text '{text}'")
            except ValueError: print(f"Warning: Could not parse color '{style['color']}' for text '{text}'")
        if 'alignment' in style: paragraph.alignment = style['alignment']
        if text.strip().startswith('- '):
            paragraph.style = 'ListBullet'
            run.text = text.strip()[2:] 
        elif re.match(r'^\d+\.\s+', text.strip()):
            paragraph.style = 'ListNumber'
            run.text = re.sub(r'^\d+\.\s+', '', text.strip())

    if table_data:
        if not table_data[0]: 
             print("Warning: Table data found but first row is empty (no columns). Skipping table.")
        else:
            try:
                num_rows = len(table_data)
                num_cols = len(table_data[0]) if num_rows > 0 else 0
                if num_rows > 0 and num_cols > 0:
                    if paragraphs_data: doc.add_paragraph() 
                    table_obj = doc.add_table(rows=num_rows, cols=num_cols)
                    table_obj.style = 'TableGrid' 
                    table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for i, row_cells in enumerate(table_data):
                        if len(row_cells) != num_cols:
                            print(f"Warning: Row {i} has {len(row_cells)} cells, expected {num_cols}. Adjusting.")
                        for j, cell_text in enumerate(row_cells):
                            if j < num_cols: 
                                table_obj.cell(i, j).text = str(cell_text)
                else:
                    print("Warning: Table data was present but resulted in 0 rows or 0 columns.")
            except Exception as e:
                print(f"Error creating table in DOCX: {e}")
    doc.save(output_path)

def process_format_request_and_save_docx(message, extracted_pdf_text, format_caption_from_user):
    bot.reply_to(message, "تم استلام طلب التنسيق. جارٍ معالجة النص مع الذكاء الاصطناعي لإنشاء ملف Word...")

    ai_formatted_text = process_text_with_openrouter_for_docx(extracted_pdf_text, format_caption_from_user)
    if "Error" in ai_formatted_text and any(err_kw in ai_formatted_text for err_kw in ["AI", "model", "OpenRouter", "API"]):
        bot.reply_to(message, f"عذرًا، حدث خطأ أثناء التواصل مع الذكاء الاصطناعي لتنسيق النص: {ai_formatted_text}")
        return

    if not ai_formatted_text.strip():
        bot.reply_to(message, "عذرًا، لم يتمكن الذكاء الاصطناعي من إنشاء محتوى منسق. حاول مرة أخرى أو بطلب مختلف.")
        return

    cleaned_ai_text_for_parsing = clean_markdown_for_docx(ai_formatted_text)
    paragraphs_data, table_data = parse_formatted_text_for_docx(cleaned_ai_text_for_parsing)

    if not paragraphs_data and not table_data:
        bot.reply_to(message, "لم أتمكن من تحليل النص المنسق من الذكاء الاصطناعي لإنشاء ملف Word. قد يكون التنسيق معقدًا جدًا أو غير متوقع.")
        return

    temp_docx_path = os.path.join(tempfile.gettempdir(), f"output_{message.chat.id}.docx")
    try:
        save_to_docx(paragraphs_data, table_data, temp_docx_path)
    except Exception as e:
        bot.reply_to(message, f"حدث خطأ أثناء إنشاء ملف Word: {e}")
        print(f"Error in save_to_docx: {e}")
        return

    try:
        with open(temp_docx_path, "rb") as docx_file:
            bot.send_document(message.chat.id, docx_file, caption="تم تجهيز الملف بتنسيق Word حسب طلبك.")
    except FileNotFoundError:
        bot.reply_to(message, "عذرًا، لم أتمكن من العثور على الملف المنشأ لإرساله.")
    except Exception as e:
        bot.reply_to(message, f"حدث خطأ أثناء إرسال ملف Word: {e}")
    finally:
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)

# ==============================================================================
# Plotting and Mathematical Functions
# ==============================================================================
def extract_plot_code(text_from_ai):
    match = re.search(r"```(?:python)?\s*([\s\S]*?)\s*```", text_from_ai)
    if match:
        return match.group(1).strip()
    return ""

def execute_plot_code(plot_code_str):
    if not plot_code_str: return None
    modified_code = re.sub(r"plt\.show\(\s*\)", "", plot_code_str)
    modified_code += "\nimport io\nplot_buffer = io.BytesIO()\nplt.savefig(plot_buffer, format='png', bbox_inches='tight')\nplt.close('all')\nplot_buffer.seek(0)\n"
    local_vars = {}
    try:
        exec_globals = {'plt': plt, 'np': np, 'sp': sp, 'io': io}
        exec(modified_code, exec_globals, local_vars)
        return local_vars.get("plot_buffer", None)
    except Exception as e:
        print(f"Error executing plot code: {e}\nCode:\n{modified_code}")
        return None

def solve_and_plot_expression_sympy(expression_str):
    try:
        plt.close('all') 
        x = sp.symbols('x')
        expr = sp.sympify(expression_str) 

        if expr.has(x):
            f = sp.lambdify(x, expr, modules=['numpy'])
            x_vals = np.linspace(-10, 10, 400) 
            y_vals = np.empty_like(x_vals)
            valid_indices = []
            for i, val in enumerate(x_vals):
                try:
                    y_val = f(val)
                    if np.iscomplex(y_val) or np.isnan(y_val) or np.isinf(y_val):
                        y_vals[i] = np.nan 
                    else:
                        y_vals[i] = y_val
                        valid_indices.append(i)
                except (ValueError, TypeError, ZeroDivisionError):
                    y_vals[i] = np.nan

            if not valid_indices: 
                print(f"Sympy could not plot expression '{expression_str}' (no valid points).")
                return None

            plt.figure(figsize=(8, 6))
            plt.plot(x_vals[valid_indices], y_vals[valid_indices], label=f"y = {sp.latex(expr)}")
            plt.xlabel("x")
            plt.ylabel("y")
            plt.title(f"الرسم البياني للدالة: ${sp.latex(expr)}$")
            plt.legend()
            plt.grid(True)
            plt.axhline(0, color='black', linewidth=0.5)
            plt.axvline(0, color='black', linewidth=0.5)

            buf = io.BytesIO()
            plt.savefig(buf, format='png', bbox_inches='tight')
            buf.seek(0)
            plt.close()
            return buf
        else: 
            print(f"Sympy: Expression '{expression_str}' does not contain 'x' or is not directly plottable as f(x).")
            return None
    except (sp.SympifyError, TypeError, Exception) as e:
        print(f"Error in solve_and_plot_expression_sympy for '{expression_str}': {e}")
        return None

# ==============================================================================
# Telegram Bot Handlers
# ==============================================================================

@bot.message_handler(commands=['start'])
def send_welcome(message):
    user_id = message.from_user.id
    username = message.from_user.username or f"User_{user_id}"
    add_user(user_id, username) 

    welcome_text = (
        "مرحباً بك! أنا بوت متعدد المهام.\n"
        "يمكنني مساعدتك في:\n"
        "- استخراج النصوص من ملفات PDF وتنسيقها في ملفات Word.\n"
        "- حل المسائل الرياضية وتقديم شروحات ورسوم بيانية.\n"
        "- الدردشة العامة والإجابة على أسئلتك.\n\n"
        "كيف يمكنني خدمتك اليوم؟\n\n"
        "للبدء بتنسيق PDF: أرسل لي ملف PDF ثم اتبع التعليمات.\n"
        "لحل مسألة: استخدم الأمر /solve ثم اكتب مسألتك."
    )
    bot.reply_to(message, welcome_text)
    if user_id in active_sessions:
        active_sessions[user_id].pop("status", None)
        active_sessions[user_id].pop("extracted_text", None)
        active_sessions[user_id].pop("chat_history", None) # Clear chat history too


@bot.message_handler(commands=['solve'])
def handle_solve_command(message):
    if not model_ready:
        bot.reply_to(message, "عذرًا، خدمة الذكاء الاصطناعي غير متاحة حاليًا. يرجى المحاولة لاحقًا.")
        return

    user_id = message.from_user.id
    problem_description_full = message.text.strip()
    command_parts = problem_description_full.split(maxsplit=1)

    if len(command_parts) < 2 or not command_parts[1].strip():
        bot.reply_to(message, "يرجى إدخال وصف المسألة أو التعبير الرياضي بعد الأمر /solve.\nمثال: `/solve مساحة المثلث للدالة y = x^2 بين x=0 و x=2`")
        return

    problem_description = command_parts[1].strip()
    bot.reply_to(message, f"تم استلام مسألتك: \"{problem_description}\"\nجارٍ التفكير في الحل... 🤔")

    session_data = active_sessions.setdefault(user_id, {})
    chat_history = session_data.get("chat_history")
    if not chat_history:
        chat_history = [
            {"role": "system", "content": "أنت مساعد خبير في الرياضيات والهندسة والفيزياء. قدم حلولاً مفصلة وواضحة للمسائل. إذا كانت المسألة تتضمن رسمًا بيانيًا يمكن إنشاؤه باستخدام مكتبة matplotlib في بايثون، قم بتضمين كود بايثون لإنشاء هذا الرسم في نهاية إجابتك، داخل بلوك كود محاط بـ ```python ... ```. تأكد أن الكود يستخدم متغيرات واضحة ويولد رسمًا دقيقًا للمسألة."},
            {"role": "user", "content": f"المسألة:\n{problem_description}"}
        ]
    else: # Append to existing history
        chat_history.append({"role": "user", "content": f"المسألة:\n{problem_description}"})

    session_data["chat_history"] = chat_history # Store updated history

    ai_response_text = send_message_to_openrouter(chat_history, GENERATION_CONFIG)

    if "Error" in ai_response_text and any(err_kw in ai_response_text for err_kw in ["AI", "model", "OpenRouter", "API"]):
        bot.reply_to(message, f"عذرًا، حدث خطأ أثناء معالجة طلبك مع الذكاء الاصطناعي: {ai_response_text}")
        # Optionally remove the last user message from history if AI failed
        if chat_history and chat_history[-1]["role"] == "user":
             chat_history.pop()
        return

    # Add AI response to history
    chat_history.append({"role": "assistant", "content": ai_response_text})
    session_data["chat_history"] = chat_history

    explanation_full = ai_response_text
    plot_code = extract_plot_code(explanation_full)
    explanation_without_code = re.sub(r"```(?:python)?\s*[\s\S]*?\s*```", "", explanation_full).strip()
    explanation_without_code = re.sub(r"إليك كود بايثون.*?:", "", explanation_without_code, flags=re.IGNORECASE).strip()
    explanation_without_code = re.sub(r"كود بايثون للرسم.*?:", "", explanation_without_code, flags=re.IGNORECASE).strip()

    plot_image_buffer = None
    if plot_code:
        print(f"Extracted plot code for /solve:\n{plot_code}")
        plot_image_buffer = execute_plot_code(plot_code)

    if not plot_image_buffer and not plot_code: # If AI didn't provide code, try direct Sympy plot
        print(f"No plot code from AI for '{problem_description}', trying Sympy direct plot.")
        plot_image_buffer = solve_and_plot_expression_sympy(problem_description)

    if explanation_without_code:
        bot.reply_to(message, f"الحل والشرح:\n\n{explanation_without_code}")
    else:
        bot.reply_to(message, "لم أتمكن من إنشاء شرح نصي. قد تكون المسألة معقدة جدًا أو غير واضحة.")

    if plot_image_buffer:
        try:
            plot_image_buffer.seek(0) 
            bot.send_photo(message.chat.id, plot_image_buffer, caption=f"الرسم البياني المتعلق بـ: {problem_description}")
        except Exception as e:
            print(f"Error sending plot photo: {e}")
            bot.send_message(message.chat.id, "تم إنشاء الرسم البياني ولكن حدث خطأ أثناء إرساله.")
        finally:
            plot_image_buffer.close()
    elif plot_code and not plot_image_buffer : 
       bot.send_message(message.chat.id, "حاولت إنشاء رسم بياني بناءً على الكود المستخرج، ولكن فشل التنفيذ.")


@bot.message_handler(content_types=['document'])
def handle_document_upload(message):
    user_id = message.from_user.id
    doc = message.document

    if not doc.file_name.lower().endswith('.pdf'):
        bot.reply_to(message, "يرجى إرسال ملف بصيغة PDF فقط.")
        return

    bot.reply_to(message, f"تم استلام ملف PDF: \"{doc.file_name}\". جارٍ تحميل واستخراج النص...")

    try:
        file_info = bot.get_file(doc.file_id)
        downloaded_file_bytes = bot.download_file(file_info.file_path)
    except Exception as e:
        bot.reply_to(message, f"عذرًا، حدث خطأ أثناء تحميل الملف: {e}")
        print(f"Error downloading file: {e}")
        return

    temp_dir = tempfile.gettempdir()
    safe_filename = "".join(c if c.isalnum() or c in ('.', '_') else '_' for c in doc.file_name)
    temp_pdf_path = os.path.join(temp_dir, f"{user_id}_{safe_filename}")

    try:
        with open(temp_pdf_path, 'wb') as new_file:
            new_file.write(downloaded_file_bytes)
    except Exception as e:
        bot.reply_to(message, f"عذرًا، حدث خطأ أثناء حفظ الملف مؤقتًا: {e}")
        print(f"Error saving temp PDF: {e}")
        return

    extracted_text = extract_text_from_pdf(temp_pdf_path)

    if os.path.exists(temp_pdf_path): 
        try:
            os.remove(temp_pdf_path)
        except Exception as e:
            print(f"Error removing temp PDF {temp_pdf_path}: {e}")

    if not extracted_text:
        bot.reply_to(message, "لم أتمكن من استخراج أي نص من ملف PDF. قد يكون الملف فارغًا، صورة بالكامل بدون طبقة نصية قابلة للاستخراج، أو محميًا بشكل قوي. حاول استخدام ملف آخر أو تأكد أن الملف يحتوي على نصوص.")
        return

    active_sessions.setdefault(user_id, {})
    active_sessions[user_id]["extracted_text"] = extracted_text
    active_sessions[user_id]["status"] = "awaiting_docx_format_prompt"
    active_sessions[user_id].pop("chat_history", None) # Clear previous chat history for new task

    bot.reply_to(message, "تم استخراج النص بنجاح! ✅\nالآن، يرجى إرسال وصف لكيفية تنسيق هذا النص في ملف Word.\n"
                          "مثال: \"اجعل العناوين الرئيسية بالخط العريض وحجم 16، ولون النص الأساسي أزرق داكن. أنشئ جدولاً بالبيانات التالية...\"")


@bot.message_handler(func=lambda message: active_sessions.get(message.from_user.id, {}).get("status") == "awaiting_docx_format_prompt")
def handle_docx_format_prompt(message):
    user_id = message.from_user.id
    session_data = active_sessions.get(user_id) 

    if not session_data or "extracted_text" not in session_data:
        bot.reply_to(message, "حدث خطأ ما، لا يوجد نص مستخرج مرتبط بجلستك. يرجى إعادة إرسال ملف PDF أولاً.")
        if session_data: session_data.pop("status", None) 
        return

    format_caption = message.text.strip()
    extracted_text = session_data["extracted_text"]

    if not format_caption:
        bot.reply_to(message, "لم تقدم وصفًا للتنسيق. يرجى توضيح كيف تريد تنسيق الملف.")
        return 

    process_format_request_and_save_docx(message, extracted_text, format_caption)

    session_data.pop("extracted_text", None)
    session_data.pop("status", None)
    if not session_data: 
        active_sessions.pop(user_id, None)


@bot.message_handler(func=lambda message: True) 
def handle_general_text_messages(message):
    if not model_ready:
        bot.reply_to(message, "عذرًا، خدمة الذكاء الاصطناعي غير متاحة حاليًا. يرجى المحاولة لاحقًا.")
        return

    user_id = message.from_user.id
    username = message.from_user.username or f"User_{user_id}" 
    user_message_text = message.text.strip()
    user_message_lower = user_message_text.lower()

    update_user_username_if_changed(user_id, username)

    identified_as = user_names_identified.get(user_id)
    if not identified_as: 
        if any(alias.lower() in user_message_lower for alias in yusuf_aliases) or user_id == 5770755631: 
            identified_as = "يوسف"
            user_names_identified[user_id] = identified_as
        elif any(alias.lower() in user_message_lower for alias in assim_aliases) or user_id == 797691024: 
            identified_as = "عاصم"
            user_names_identified[user_id] = identified_as
        elif any(alias.lower() in user_message_lower for alias in basmala_aliases) or user_id == 5708090246: 
            identified_as = "بسملة"
            user_names_identified[user_id] = identified_as

    if "من أنا" in user_message_text or "انا مين" in user_message_text:
        if identified_as:
            bot.reply_to(message, f"أنت {identified_as} بالتأكيد! صديق عزيز جدًا. 😊")
        else:
            bot.reply_to(message, "أنت المستخدم الذي أتحدث معه الآن. لم أتعرف على اسمك بعد، هل يمكنك إخباري به؟")
        return

    if any(greeting.lower() in user_message_lower for greeting in greeting_words):
        if identified_as:
            bot.reply_to(message, f"أهلاً بك يا {identified_as}! سعيد جدًا بالتحدث معك. كيف يمكنني مساعدتك اليوم؟")
        else:
            bot.reply_to(message, "مرحباً بك! كيف يمكنني مساعدتك؟ إذا أردت، يمكنك إخباري باسمك.")
        return

    if user_message_text.startswith("انا ") or user_message_text.startswith("أنا "):
        claimed_name_part = user_message_text.split(maxsplit=1)[1]
        if any(alias.lower() in claimed_name_part.lower() for alias in yusuf_aliases):
            if user_id == 5770755631: 
                bot.reply_to(message, "أهلاً بك يا يوسف! كيف يمكنني مساعدتك؟")
                if not identified_as: user_names_identified[user_id] = "يوسف"
            else:
                bot.reply_to(message, f"أنت لست يوسف الذي أعرفه. حساب يوسف الحقيقي هو @youssz (مثال).")
            return
        elif any(alias.lower() in claimed_name_part.lower() for alias in assim_aliases):
            if user_id == 797691024: 
                bot.reply_to(message, "أهلاً بك يا عاصم! كيف يمكنني مساعدتك؟")
                if not identified_as: user_names_identified[user_id] = "عاصم"
            else:
                bot.reply_to(message, f"أنت لست عاصم الذي أعرفه. حساب عاصم الحقيقي هو @AssemShimi (مثال).")
            return
        elif any(alias.lower() in claimed_name_part.lower() for alias in basmala_aliases):
            if user_id == 5708090246: 
                bot.reply_to(message, "أهلاً بكِ يا بسملة! كيف يمكنني مساعدتكِ؟")
                if not identified_as: user_names_identified[user_id] = "بسملة"
            else:
                bot.reply_to(message, f"أنتِ لستِ بسملة التي أعرفها. حساب بسملة الحقيقي هو @BASMALA_MABROUK (مثال).")
            return

    if "مميز" in user_message_text or "فائق" in user_message_text:
        special_status = is_user_special(user_id)
        if "فائق" in user_message_text: 
            if special_status == 2:
                bot.reply_to(message, "أنت بالفعل مستخدم فائق! 💎 استمتع بالميزات الكاملة.")
            else:
                bot.reply_to(message, "للحصول على ترقية إلى مستخدم فائق، يرجى التواصل مع المطور @AssemShimi (مثال).")
        elif "مميز" in user_message_text:
            if special_status == 1:
                bot.reply_to(message, "أنت بالفعل مستخدم مميز! ✨ لديك صلاحيات إضافية.")
            elif special_status == 2: 
                bot.reply_to(message, "أنت مستخدم فائق، وهذا يشمل ميزات المستخدم المميز وأكثر! 💎")
            else:
                bot.reply_to(message, "للحصول على ترقية إلى مستخدم مميز، يرجى التواصل مع المطور @AssemShimi (مثال).")
        return

    session_data = active_sessions.setdefault(user_id, {})
    chat_history = session_data.get("chat_history")
    if not chat_history:
        initial_chat_history = [
            {"role": "system", "content": "أنا صوما، نموذج ذكاء اصطناعي. أنا هنا لمساعدتك في مهام متنوعة مثل الإجابة على الأسئلة، معالجة المستندات، وغيرها الكثير."},
            {"role": "user", "content": "Hi"},
            {"role": "assistant", "content": "Hi there! I am Sooma, an AI assistant. How can I help you today?"},
            {"role": "user", "content": "Who are you?"},
            {"role": "assistant", "content": "I am Sooma, an AI model. I can help with various tasks like answering questions, processing documents, and more."},
             {"role": "user", "content": "من أنت؟"},
            {"role": "assistant", "content": "أنا صوما، نموذج ذكاء اصطناعي. أنا هنا لمساعدتك في مهام متنوعة مثل الإجابة على الأسئلة، معالجة المستندات، وغيرها الكثير."}
        ]
        if identified_as:
            initial_chat_history.append({"role": "user", "content": f"My name is {identified_as}."})
            initial_chat_history.append({"role": "assistant", "content": f"Nice to meet you, {identified_as}!"})
        chat_history = initial_chat_history

    # Add current user message to history
    chat_history.append({"role": "user", "content": user_message_text})

    ai_reply_text = send_message_to_openrouter(chat_history, GENERATION_CONFIG)

    if "Error" not in ai_reply_text: # Only add assistant response if no error
       chat_history.append({"role": "assistant", "content": ai_reply_text})
    else: # If error, remove the last user message to avoid resending it in a broken state
        if chat_history and chat_history[-1]["role"] == "user":
            chat_history.pop()

    session_data["chat_history"] = chat_history # Save updated history

    user_db_status = is_user_special(user_id)
    response_prefix = ""
    if user_db_status == 2: response_prefix = "مستخدم فائق 💎\n\n"
    elif user_db_status == 1: response_prefix = "مستخدم مميز ✨\n\n"

    bot.reply_to(message, response_prefix + ai_reply_text)

# ==============================================================================
# Admin Console Functions (Not Telegram commands)
# ==============================================================================
def show_users_in_console():
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute('SELECT user_id, username, is_special FROM users')
        users = cursor.fetchall()
        conn.close()
        if users:
            print("\n📋 قائمة المستخدمين المسجلين:")
            for uid, uname, special_stat in users:
                status_text = {0: "عادي", 1: "مميز", 2: "فائق"}.get(special_stat, "غير معروف")
                print(f"- ID: {uid}, Username: @{uname if uname else 'N/A'}, Status: {status_text} ({special_stat})")
        else:
            print("⚠️ لا يوجد مستخدمون في قاعدة البيانات.")
    except sqlite3.Error as e:
        print(f"❌ حدث خطأ أثناء الوصول إلى قاعدة البيانات لعرض المستخدمين: {e}")

def update_special_status_in_db(target_user_id, new_status_code):
    if new_status_code not in [0, 1, 2]:
        print(f"❌ كود الحالة '{new_status_code}' غير صالح. استخدم 0, 1, أو 2.")
        return
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute('UPDATE users SET is_special = ? WHERE user_id = ?', (new_status_code, target_user_id))
        conn.commit()
        if cursor.rowcount > 0:
            status_text = {0: "عادي", 1: "مميز", 2: "فائق"}.get(new_status_code)
            print(f"✅ تم تحديث حالة المستخدم {target_user_id} إلى '{status_text}' ({new_status_code}) بنجاح.")
        else:
            print(f"⚠️ لم يتم العثور على المستخدم بالمعرف {target_user_id} لتحديث حالته.")
    except sqlite3.Error as e:
        print(f"❌ حدث خطأ أثناء تحديث حالة المستخدم: {e}")
    finally:
        if conn:
            conn.close()

# ==============================================================================
# Main Bot Execution
# ==============================================================================
if __name__ == '__main__':
    print("Setting up database...")
    setup_database()
    show_users_in_console() 

    if not model_ready or TELEGRAM_TOKEN == "YOUR_TELEGRAM_BOT_TOKEN":
        print("\nCritical Error: API keys not set or AI model failed to initialize properly.")
        print("Please configure OPENROUTER_API_KEY and TELEGRAM_TOKEN in the script.")
        if not model_ready:
            print("OPENROUTER_API_KEY might be missing, incorrect, or the model name is invalid for OpenRouter.")
        if TELEGRAM_TOKEN == "YOUR_TELEGRAM_BOT_TOKEN":
            print("TELEGRAM_TOKEN is not set.")
        print("The bot will not start.\n")
    else:
        print(f"\nBot '{bot.get_me().username}' is starting with OpenRouter model '{OPENROUTER_MODEL_NAME}'...")
        print("Listening for messages...")
        try:
            bot.infinity_polling(skip_pending=True) 
        except Exception as main_loop_error:
            print(f"An error occurred in the main bot loop: {main_loop_error}")
            print("The bot might have stopped. Check logs and configurations.")
